"""
SOW Report Generator - Creates professional Word document reports

Generates Statement of Work (SOW) reports with:
1. Scope of Service
2. Timings & Effort Estimation  
3. Key Design Decisions (KDD)
"""

from datetime import datetime, timedelta
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise ImportError("python-docx is required. Install with: pip install python-docx")

from backend.config import OUTPUT_DIR
from backend.data.excel_templates import KDD_DEFINITIONS, KDD_CONDITIONAL_MAPPINGS


class SOWReportGenerator:
    """Generate professional SOW report as Word document"""
    
    def __init__(self):
        """Initialize report generator"""
        pass  # No Excel loading needed anymore
    
    def get_scope_details(self, scope_inputs_dict):
        """Extract scope details from user input (scope_inputs_dict)
        
        Args:
            scope_inputs_dict: Dict mapping metric names to {'in_scope': str, 'details': number}
        
        Returns:
            Dict with actual user input values
        """
        details = {
            'Account': scope_inputs_dict.get('Account', {}).get('details', 0),
            'Account Hierarchies': scope_inputs_dict.get('Account Alternate Hierarchies', {}).get('details', 0),
            'Entity': scope_inputs_dict.get('Entity', {}).get('details', 0),
            'Entity Hierarchies': scope_inputs_dict.get('Entity Alternate Hierarchies', {}).get('details', 0),
            'Currencies': scope_inputs_dict.get('Multi-Currency', {}).get('details', 0),
            'Reporting Currencies': scope_inputs_dict.get('Reporting Currency', {}).get('details', 0),
            'Custom Dimensions': scope_inputs_dict.get('Custom Dimensions', {}).get('details', 0),
            'Custom Dimension Hierarchies': scope_inputs_dict.get('Alternate Hierarchies in Custom Dimensions', {}).get('details', 0),
            'Data Forms': scope_inputs_dict.get('Data Forms', {}).get('details', 0),
            'Business Rules': scope_inputs_dict.get('Business Rules', {}).get('details', 0),
            'Member Formulas': scope_inputs_dict.get('Member Formula', {}).get('details', 0),
        }
        return details
    
    def get_applicable_kdds(self, scope_metrics: list):
        """
        Get KDD items based on scope definition values
        
        Returns list of (kdd_id, kdd_text) tuples
        First 4 are always included, rest are conditional based on scope metrics
        """
        # First 4 KDD items - Always included
        default_kdds = [
            ('KDD01', 'General Application Configuration'),
            ('KDD02', 'Metadata Configuration'),
            ('KDD03', 'FCC Consolidations and Other Calculations'),
            ('KDD04', 'Reports and Data Form Configuration'),
        ]
        
        kdd_items = list(default_kdds)
        
        # Create a lookup dict of metric name -> in_scope flag from scope_metrics
        scope_dict = {m['name']: m['in_scope_flag'] for m in scope_metrics}
        
        # Check conditional KDD mappings from template
        for scope_metric_name, kdd_key, _ in KDD_CONDITIONAL_MAPPINGS:
            # Check if this metric is in scope (in_scope_flag = 1)
            if scope_dict.get(scope_metric_name, 0) > 0:
                kdd_text = KDD_DEFINITIONS.get(kdd_key, '')
                if kdd_text:
                    kdd_items.append((kdd_key, kdd_text))
        
        return kdd_items
    
    def generate_word_document(self, scope_result, effort_estimation, summary, fte_allocation, scope_inputs_dict, output_path=None):
        """
        Generate professional SOW report as Word document
        
        Args:
            scope_result: Output from ScopeDefinitionProcessor
            effort_estimation: Output from EffortCalculator
            summary: Summary dict with total hours/days/months
            fte_allocation: Dict of role -> {'hours': value, ...}
            scope_inputs_dict: Dict with user input scope values
            output_path: Path to save the Word document
        
        Returns:
            Path to generated document
        """
        # Create Document
        doc = Document()
        
        # Get data
        scope_data = self.get_scope_details(scope_inputs_dict)
        weightage = scope_result['total_weightage']
        tier_name = scope_result['tier_name']
        tier_range = scope_result['tier_range']
        scope_metrics = scope_result.get('metrics', [])
        
        # Calculate dates
        start_date = datetime.now()
        num_months = round(summary['total_months'])
        end_date = start_date + timedelta(days=num_months * 30)
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # ===== TITLE PAGE =====
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('STATEMENT OF WORK (SOW)')
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run('FCC IMPLEMENTATION ENGAGEMENT')
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.bold = True
        
        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Generated: {start_date.strftime('%d-%b-%Y')}").font.size = Pt(10)
        
        doc.add_paragraph(f"Engagement Weightage: {weightage:.1f}", style='List Bullet')
        doc.add_paragraph(f"Implementation Tier: {tier_name} ({tier_range[0]}-{tier_range[1]})", style='List Bullet')
        
        doc.add_paragraph()  # Spacing
        
        # ===== SECTION 1: SCOPE OF SERVICE =====
        doc.add_heading('1. SCOPE OF SERVICE', level=1)
        
        doc.add_paragraph(f"Engagement Tier: {tier_name}")
        
        doc.add_paragraph(
            "Under this SOW, Donyati will work with Client to provide Services noted below. "
            "Actual activities, work items, schedule and deliverables would be jointly managed "
            "by Donyati and Client based on the Client's business priorities."
        )
        
        doc.add_paragraph(
            "The scope of this engagement is focused on the implementation of the following "
            "application modules: Oracle EPM Enterprise – Financial Consolidation and Close"
        )
        
        # Dimensions
        doc.add_heading('DIMENSIONS', level=2)
        doc.add_paragraph(
            f"Account Dimension: Approximately {scope_data.get('Account', 'TBD')} accounts will be "
            "configured based on the current Chart of Accounts.",
            style='List Bullet'
        )
        doc.add_paragraph(
            f"Account Alternate Hierarchies: Up to {scope_data.get('Account Hierarchies', 'TBD')} "
            "alternate hierarchies will be developed.",
            style='List Bullet'
        )
        doc.add_paragraph(
            f"Entity Dimension: Approximately {scope_data.get('Entity', 'TBD')} entities will be "
            "configured based on the current structure.",
            style='List Bullet'
        )
        doc.add_paragraph(
            f"Entity Alternate Hierarchies: Up to {scope_data.get('Entity Hierarchies', 'TBD')} "
            "alternate hierarchies will be developed.",
            style='List Bullet'
        )
        doc.add_paragraph(
            f"Currency Configuration: {scope_data.get('Currencies', 'TBD')} currencies will be "
            f"configured with {scope_data.get('Reporting Currencies', 'TBD')} reporting currency/currencies.",
            style='List Bullet'
        )
        doc.add_paragraph(
            f"Custom Dimensions: {scope_data.get('Custom Dimensions', 'TBD')} custom dimensions will be "
            f"leveraged to support additional reporting requirements. Up to "
            f"{scope_data.get('Custom Dimension Hierarchies', 'TBD')} alternate hierarchies will be developed.",
            style='List Bullet'
        )
        doc.add_paragraph(
            "Standard Dimensions: Year, Period, View, Consolidation, Intercompany, and Data Source "
            "dimensions will be configured to support consolidation and reporting.",
            style='List Bullet'
        )
        
        # Application Features
        doc.add_heading('APPLICATION FEATURES', level=2)
        doc.add_paragraph("Standard elimination capabilities will be provided", style='List Bullet')
        doc.add_paragraph("Consolidation journals will be enabled", style='List Bullet')
        doc.add_paragraph("Consolidation journal templates will be created as needed", style='List Bullet')
        doc.add_paragraph("Approval process will be utilized in the application", style='List Bullet')
        doc.add_paragraph("Task manager may be used to support Financial Consolidation and Close", style='List Bullet')
        
        # Application Customization
        doc.add_heading('APPLICATION CUSTOMIZATION', level=2)
        doc.add_paragraph(
            f"Custom Data Forms: If required, up to {scope_data.get('Data Forms', 'TBD')} custom data forms "
            "will be developed to support Financial Consolidation and Close.",
            style='List Bullet'
        )
        
        # Calculations
        doc.add_heading('CALCULATIONS', level=2)
        doc.add_paragraph(
            f"Custom Business Rules: If required, up to {scope_data.get('Business Rules', 'TBD')} custom "
            "business rules will be developed to support Financial Consolidation and Close.",
            style='List Bullet'
        )
        doc.add_paragraph(
            f"Member Formulas: If required, up to {scope_data.get('Member Formulas', 'TBD')} member formulas "
            "will be developed to support Financial Consolidation and Close.",
            style='List Bullet'
        )
        
        # ===== SECTION 2: TIMINGS =====
        doc.add_page_break()
        doc.add_heading('2. TIMINGS (EFFORT ESTIMATION)', level=1)
        
        doc.add_paragraph(f"The engagement start date is {start_date.strftime('%d-%b-%Y')}")
        doc.add_paragraph(f"The engagement end date is {end_date.strftime('%d-%b-%Y')}")
        doc.add_paragraph(
            "Go-live support will be provided and additional capabilities will be added per a "
            "subsequent Statement of Work expected to be started immediately following the "
            "completion of this Statement of Work."
        )
        doc.add_paragraph(f"This SOW assumes {num_months} months to complete Financial Consolidation and Close.")
        
        doc.add_paragraph(
            "Completion of Services and Deliverables agreed upon is subject to, among other things, "
            "appropriate cooperation, obtaining the necessary information, and timely response to inquiries. "
            "The chart below shows the estimated hours per month for each resource role. At the conclusion of "
            "Requirements and Design, Donyati will revise the implementation timeline to incorporate agreed upon "
            "requirements and design elements. Donyati and Client will review and approve the revised implementation "
            "timeline before moving into the Development Phase."
        )
        
        # Create resource allocation table
        doc.add_heading('Resource Role/Monthly Hours', level=2)
        
        # Get roles in order
        roles = sorted(fte_allocation.keys())
        
        # Create table: rows = roles + 1 header, cols = months + 1 role name
        table = doc.add_table(rows=len(roles) + 1, cols=num_months + 1)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Resource Role'
        for month in range(1, num_months + 1):
            header_cells[month].text = str(month)
        
        # Data rows
        for idx, role in enumerate(roles):
            row_cells = table.rows[idx + 1].cells
            row_cells[0].text = role
            
            # Calculate monthly hours
            total_hours = fte_allocation[role]['hours']
            monthly_hours = round(total_hours / num_months)
            
            # Fill all months with same value
            for month in range(1, num_months + 1):
                row_cells[month].text = str(monthly_hours)
        
        doc.add_paragraph()  # Spacing
        
        # Effort by category
        doc.add_heading('TOTAL IMPLEMENTATION EFFORT', level=2)
        
        doc.add_paragraph(f"Total Hours: {summary['total_time_hours']:.1f} hours", style='List Bullet')
        doc.add_paragraph(f"Total Days: {summary['total_days']:.1f} days (@ 8 hours/day)", style='List Bullet')
        doc.add_paragraph(f"Total Months: {summary['total_months']:.2f} months (@ 30 days/month)", style='List Bullet')
        
        doc.add_heading('EFFORT BY CATEGORY', level=2)
        
        effort_table = doc.add_table(rows=len(effort_estimation) + 2, cols=3)
        effort_table.style = 'Light Grid Accent 1'
        
        # Header
        hdr_cells = effort_table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Hours'
        hdr_cells[2].text = 'Days'
        
        # Data
        row_idx = 1
        for category_name in sorted(effort_estimation.keys()):
            cat_data = effort_estimation[category_name]
            hours = cat_data['final_estimate']
            days = cat_data['in_days']
            
            row_cells = effort_table.rows[row_idx].cells
            row_cells[0].text = category_name
            row_cells[1].text = f"{hours:.1f}"
            row_cells[2].text = f"{days:.2f}"
            row_idx += 1
        
        # Total row
        total_cells = effort_table.rows[row_idx].cells
        total_cells[0].text = 'TOTAL'
        total_cells[1].text = f"{summary['total_time_hours']:.1f}"
        total_cells[2].text = f"{summary['total_days']:.2f}"
        
        # Make total row bold
        for cell in total_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # ===== SECTION 3: KEY DESIGN DECISIONS =====
        doc.add_page_break()
        doc.add_heading('3. KEY DESIGN DECISIONS (KDD)', level=1)
        
        doc.add_paragraph(
            "The following Key Design Decisions have been identified as critical for the "
            "successful implementation:"
        )
        
        # Get KDD items
        kdd_items = self.get_applicable_kdds(scope_metrics)
        
        if kdd_items:
            for kdd_id, kdd_text in kdd_items:
                doc.add_paragraph(f"{kdd_id}: {kdd_text}", style='List Number')
        else:
            doc.add_paragraph("No key design decisions identified for this scope.", style='Normal')
        
        # Save document
        if not output_path:
            OUTPUT_DIR.mkdir(exist_ok=True)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = OUTPUT_DIR / f'SOW_Report_{timestamp}.docx'
        
        output_path = Path(output_path)
        doc.save(str(output_path))
        
        return str(output_path)
